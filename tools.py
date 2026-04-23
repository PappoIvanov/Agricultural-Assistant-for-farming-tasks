import os
import requests
from datetime import datetime, date, timedelta
from pathlib import Path
from supabase import create_client as _create_supabase
from config import PARCELS, ROSE_CONFIG, DIARY_PATH, AGRO_DIARY_PATH


def _supabase():
    return _create_supabase(
        os.getenv("SUPABASE_URL"),
        os.getenv("SUPABASE_KEY"),
    )


LITERATURE_DIR = Path("05_Литература")
DIAGNOSTIC_DIARY_PATH = Path("05_Литература/Диагностичен_дневник.md")


def get_weather(parcel_name: str, target_date: str = None) -> dict:
    """Връща метео данни от Open-Meteo за даден парцел и дата."""
    parcel = PARCELS.get(parcel_name)
    if not parcel or not parcel["lat"]:
        return {"error": f"Няма координати за {parcel_name}"}

    if target_date is None:
        target_date = date.today().isoformat()

    url = "https://api.open-meteo.com/v1/forecast"
    params = {
        "latitude": parcel["lat"],
        "longitude": parcel["lon"],
        "daily": "temperature_2m_max,temperature_2m_min,precipitation_sum,windspeed_10m_max",
        "forecast_days": 7,
        "timezone": "Europe/Sofia",
    }
    try:
        r = requests.get(url, params=params, timeout=10)
        response_json = r.json()
        if "daily" not in response_json:
            return {"error": "Метео услугата не отговори. Не пръскай без да знаеш условията."}
        data = response_json["daily"]
        dates = data["dates"] if "dates" in data else data.get("time", [])
        result = []
        for i, d in enumerate(dates):
            result.append({
                "date": d,
                "temp_max": data["temperature_2m_max"][i],
                "temp_min": data["temperature_2m_min"][i],
                "precip": data["precipitation_sum"][i],
                "wind_max": data["windspeed_10m_max"][i],
            })
        return {"parcel": parcel_name, "forecast": result}
    except Exception as e:
        return {"error": f"Метео услугата не отговори ({e}). Не пръскай без да знаеш условията."}


def save_spray_record(
    parcel: str,
    products: list,
    volume_liters: float,
    nozzle_count: int,
    notes: str = "",
    record_date: str = None,
) -> dict:
    """Записва пръскане в работния дневник (markdown)."""
    if record_date is None:
        record_date = date.today().isoformat()

    products_table = "\n".join(
        f"| {p['name']} | {p.get('dose', '—')} | {p.get('amount', '—')} | — |"
        for p in products
    )

    entry = f"""
## Запис: {record_date}

**Операция:** Пръскане
**Парцел:** {parcel}
**Работен разтвор:** {volume_liters} л
**Брой дюзи:** {nozzle_count}

### Използвани препарати
| Продукт | Доза / дка | Общо | Цена лв. |
|---|---|---|---|
{products_table}

### Наблюдения
{notes if notes else '—'}

---
"""
    try:
        diary = Path(DIARY_PATH)
        with open(diary, "a", encoding="utf-8") as f:
            f.write(entry)
        return {"status": "ok", "message": f"Записано в дневника за {record_date}"}
    except Exception as e:
        return {"error": str(e)}


def read_spray_history(parcel: str = None, limit: int = 5) -> dict:
    """Чете последните записи от работния дневник."""
    try:
        diary = Path(DIARY_PATH)
        content = diary.read_text(encoding="utf-8")
        entries = [e.strip() for e in content.split("---") if e.strip()]
        if parcel:
            entries = [e for e in entries if parcel in e]
        return {"entries": entries[-limit:]}
    except Exception as e:
        return {"error": str(e)}


def calculate_concentration(
    dose_per_dka: float,
    area_dka: float,
    volume_liters: float,
    row_spacing_m: float = None,
    canopy_width_m: float = None,
) -> dict:
    """
    Изчислява количеството препарат за работния разтвор.
    Отчита ефективната площ при редови насаждения.
    """
    row_spacing = row_spacing_m or ROSE_CONFIG.get("row_spacing_m")
    canopy_width = canopy_width_m or ROSE_CONFIG.get("canopy_width_m")

    if row_spacing and canopy_width:
        correction = canopy_width / row_spacing
        effective_area = area_dka * correction
    else:
        correction = 1.0
        effective_area = area_dka

    total_product_ml = dose_per_dka * effective_area

    return {
        "dose_per_dka_ml": dose_per_dka,
        "total_area_dka": area_dka,
        "effective_area_dka": round(effective_area, 2),
        "correction_factor": round(correction, 3),
        "total_product_ml": round(total_product_ml, 1),
        "concentration_ml_per_l": round(total_product_ml / volume_liters, 2),
        "note": "Ефективната площ е коригирана за междуредия." if correction < 1 else
                "Няма корекция — използвана е пълната площ.",
    }



def list_literature() -> dict:
    """Показва всички файлове в 05_Литература/ рекурсивно."""
    if not LITERATURE_DIR.exists():
        return {"error": "Папката 05_Литература/ не е намерена"}
    files = []
    for p in LITERATURE_DIR.rglob("*"):
        if p.is_file():
            files.append({
                "name": p.name,
                "path": str(p.relative_to(LITERATURE_DIR)),
                "type": p.suffix.lower(),
                "size_kb": round(p.stat().st_size / 1024, 1),
            })
    return {"files": files, "total": len(files)}


MAX_LITERATURE_CHARS = 12000


def _truncate(text: str) -> str:
    if len(text) > MAX_LITERATURE_CHARS:
        return text[:MAX_LITERATURE_CHARS] + f"\n\n[... файлът е съкратен до {MAX_LITERATURE_CHARS} символа]"
    return text


def read_literature(filename: str) -> dict:
    """Чете файл от 05_Литература/. Поддържа .txt, .md, .pdf, .docx."""
    matches = list(LITERATURE_DIR.rglob(filename))
    if not matches:
        return {"error": f"Файлът '{filename}' не е намерен в 05_Литература/"}

    path = matches[0]
    suffix = path.suffix.lower()

    try:
        if suffix in (".txt", ".md"):
            return {"filename": filename, "content": path.read_text(encoding="utf-8")}

        elif suffix == ".docx":
            from docx import Document
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {"filename": filename, "content": _truncate(text)}

        elif suffix == ".pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        pages.append(f"[Страница {i}]\n{text}")
            return {"filename": filename, "content": _truncate("\n\n".join(pages))}

        elif suffix in (".jpg", ".jpeg", ".png"):
            return {
                "filename": filename,
                "type": "image",
                "note": "Снимката се праща директно към vision модела — използвай upload бутона.",
            }

        else:
            return {"error": f"Неподдържан формат: {suffix}"}

    except Exception as e:
        return {"error": str(e)}


def search_literature(query: str, filename: str = None) -> dict:
    """Търси по ключови думи в литературата и връща само релевантните части."""
    def _extract_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in (".txt", ".md"):
            return path.read_text(encoding="utf-8")
        elif suffix == ".docx":
            from docx import Document
            doc = Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif suffix == ".pdf":
            import pdfplumber
            pages = []
            with pdfplumber.open(str(path)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    t = page.extract_text()
                    if t:
                        pages.append(f"[Страница {i}]\n{t}")
            return "\n\n".join(pages)
        return ""

    def _score_chunk(chunk: str, keywords: list) -> int:
        chunk_lower = chunk.lower()
        return sum(1 for kw in keywords if kw in chunk_lower)

    try:
        _SEARCH_DIRS = [LITERATURE_DIR, Path("03_Препарати_и_Торове")]
        keywords = [w.lower() for w in query.split() if len(w) > 2]
        if filename:
            files = [LITERATURE_DIR / filename]
        else:
            files = []
            for d in _SEARCH_DIRS:
                if d.exists():
                    files.extend(d.rglob("*"))
        files = [f for f in files if f.is_file() and f.suffix.lower() in (".txt", ".md", ".pdf", ".docx")]
        # Пропускай PDF/DOCX ако вече има парсната .txt или .md версия
        parsed_stems = {f.stem for f in files if f.suffix.lower() in (".txt", ".md")}
        files = [f for f in files if f.suffix.lower() in (".txt", ".md") or f.stem not in parsed_stems]

        all_results = []
        for path in files:
            text = _extract_text(path)
            chunks = [c.strip() for c in text.split("\n\n") if c.strip()]
            for chunk in chunks:
                score = _score_chunk(chunk, keywords)
                if score > 0:
                    all_results.append((score, path.name, chunk))

        all_results.sort(key=lambda x: x[0], reverse=True)
        top = all_results[:5]

        if not top:
            return {"results": [], "message": "Няма намерени резултати за: " + query}

        MAX_CHUNK_CHARS = 500
        MAX_TOTAL_CHARS = 3000
        results = []
        total = 0
        for score, name, chunk in top:
            chunk = chunk[:MAX_CHUNK_CHARS]
            if total + len(chunk) > MAX_TOTAL_CHARS:
                break
            results.append({"file": name, "score": score, "text": chunk})
            total += len(chunk)

        return {"query": query, "results": results}

    except Exception as e:
        return {"error": str(e)}


KNOWLEDGE_BASE_PATH = Path("05_Литература/База_знания.md")


def save_to_knowledge_base(title: str, content: str, source: str = None) -> dict:
    """Записва нова информация в База_знания.md — научено от снимка, литература или наблюдение."""
    today = date.today().isoformat()
    source_line = f"**Източник:** {source}" if source else ""

    entry = f"""
## {title} *(добавено {today})*

{source_line}

{content}

---
"""
    try:
        with open(KNOWLEDGE_BASE_PATH, "a", encoding="utf-8") as f:
            f.write(entry)
        return {"status": "ok", "message": f"Записано в базата знания: {title}"}
    except Exception as e:
        return {"error": str(e)}


def read_diagnostic_diary() -> dict:
    """Чете диагностичния дневник с минали случаи и корекции."""
    try:
        content = DIAGNOSTIC_DIARY_PATH.read_text(encoding="utf-8")
        return {"content": content}
    except Exception as e:
        return {"error": str(e)}


def save_diagnostic_case(
    symptoms: str,
    initial_diagnosis: str,
    correction: str = None,
    action_taken: str = None,
    outcome: str = None,
    case_date: str = None,
) -> dict:
    """Записва случай в диагностичния дневник."""
    if case_date is None:
        case_date = date.today().isoformat()

    lines = [f"\n## Случай: {case_date}"]
    lines.append(f"**Симптоми:** {symptoms}")
    lines.append(f"**Първоначална диагноза:** {initial_diagnosis}")
    if correction:
        lines.append(f"**Корекция:** {correction}")
    if action_taken:
        lines.append(f"**Предприето действие:** {action_taken}")
    if outcome:
        lines.append(f"**Резултат:** {outcome}")
    lines.append("\n---")

    entry = "\n".join(lines)
    try:
        with open(DIAGNOSTIC_DIARY_PATH, "a", encoding="utf-8") as f:
            f.write(entry)
        return {"status": "ok", "message": f"Случаят е записан за {case_date}"}
    except Exception as e:
        return {"error": str(e)}


def send_telegram(message: str) -> dict:
    """Изпраща съобщение до стопанина през Telegram."""
    token = os.getenv("TELEGRAM_BOT_TOKEN")
    chat_id = os.getenv("TELEGRAM_CHAT_ID")
    if not token or not chat_id:
        return {"error": "Липсва TELEGRAM_BOT_TOKEN или TELEGRAM_CHAT_ID в .env"}
    try:
        r = requests.post(
            f"https://api.telegram.org/bot{token}/sendMessage",
            json={"chat_id": chat_id, "text": message, "parse_mode": "HTML"},
            timeout=10,
        )
        return {"status": "ok"} if r.ok else {"error": r.text}
    except Exception as e:
        return {"error": str(e)}


def save_planned_spray(
    planned_date: str,
    parcel: str,
    products: list,
    volume_liters: float,
    nozzle_count: int,
    notes: str = "",
) -> dict:
    """Записва планирано пръскане в Supabase."""
    try:
        spray_date = date.fromisoformat(planned_date)
        if spray_date < date.today():
            return {"error": f"Датата {planned_date} е в миналото. Моля въведи бъдеща дата."}
        existing = (
            _supabase().table("planned_sprays")
            .select("id")
            .eq("planned_date", planned_date)
            .eq("parcel", parcel)
            .eq("completed", False)
            .execute()
        )
        if existing.data:
            return {"error": f"Вече има планирано пръскане за {planned_date} на {parcel}. Изтрий го първо или избери друга дата."}
        _supabase().table("planned_sprays").insert({
            "planned_date": planned_date,
            "parcel": parcel,
            "products": products,
            "volume_liters": volume_liters,
            "nozzle_count": nozzle_count,
            "notes": notes,
            "completed": False,
        }).execute()
        return {"status": "ok", "message": f"Планирано пръскане записано за {planned_date}"}
    except Exception as e:
        return {"error": str(e)}


def get_planned_sprays(days_ahead: int = 3) -> dict:
    """Връща планираните пръскания за следващите N дни от Supabase."""
    try:
        today = date.today().isoformat()
        until = (date.today() + timedelta(days=days_ahead)).isoformat()
        result = (
            _supabase().table("planned_sprays")
            .select("*")
            .eq("completed", False)
            .gte("planned_date", today)
            .lte("planned_date", until)
            .order("planned_date")
            .execute()
        )
        return {"sprays": result.data}
    except Exception as e:
        return {"error": str(e)}


def complete_planned_spray(spray_id: int) -> dict:
    """Маркира планирано пръскане като изпълнено в Supabase."""
    try:
        _supabase().table("planned_sprays").update({"completed": True}).eq("id", spray_id).execute()
        return {"status": "ok", "message": f"Пръскане #{spray_id} е маркирано като изпълнено."}
    except Exception as e:
        return {"error": str(e)}


def save_agro_operation(
    operation_type: str,
    parcel: str,
    notes: str = "",
    equipment: str = "",
    record_date: str = None,
) -> dict:
    """Записва агротехническа операция в дневника."""
    if record_date is None:
        record_date = date.today().isoformat()

    entry = f"""
## {operation_type} — {record_date}

**Парцел:** {parcel}
**Оборудване:** {equipment if equipment else "—"}
**Бележки:** {notes if notes else "—"}

---
"""
    try:
        diary = Path(AGRO_DIARY_PATH)
        diary.parent.mkdir(parents=True, exist_ok=True)
        with open(diary, "a", encoding="utf-8") as f:
            f.write(entry)
        return {"status": "ok", "message": f"{operation_type} записано за {record_date}"}
    except Exception as e:
        return {"error": str(e)}


def read_agro_history(parcel: str = None, operation_type: str = None, limit: int = 10) -> dict:
    """Чете историята на агротехническите операции."""
    try:
        diary = Path(AGRO_DIARY_PATH)
        if not diary.exists():
            return {"entries": [], "message": "Все още няма записани агротехнически операции."}
        content = diary.read_text(encoding="utf-8")
        entries = [e.strip() for e in content.split("---") if e.strip()]
        if parcel:
            entries = [e for e in entries if parcel in e]
        if operation_type:
            entries = [e for e in entries if operation_type.lower() in e.lower()]
        return {"entries": entries[-limit:]}
    except Exception as e:
        return {"error": str(e)}
