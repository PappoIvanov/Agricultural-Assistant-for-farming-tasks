"""
Преобразува PDF, DOCX и XLSX файлове към .txt или .md.
- Файлове с таблици → .md (markdown таблици)
- Файлове без таблици → .txt (обикновен текст)
- XLSX → винаги .md (таблични данни)
Пропуска файлове, които вече са обработени.
Пускай: python preprocess_literature.py
"""

from pathlib import Path

SEARCH_DIRS = [
    Path("05_Литература"),
    Path("03_Препарати_и_Торове"),
]


def _table_to_markdown(table: list) -> str:
    if not table or not table[0]:
        return ""
    rows = [[str(cell) if cell is not None else "" for cell in row] for row in table]
    header = "| " + " | ".join(rows[0]) + " |"
    separator = "| " + " | ".join("---" for _ in rows[0]) + " |"
    body = "\n".join("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n".join([header, separator, body])


def _process_pdf(path: Path) -> tuple[str, bool]:
    import pdfplumber
    pages_text = []
    has_tables = False

    with pdfplumber.open(str(path)) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            tables = page.extract_tables()
            if tables:
                has_tables = True
                page_parts = [f"## Страница {i}"]
                for table in tables:
                    page_parts.append(_table_to_markdown(table))
                plain = page.extract_text()
                if plain:
                    used_text = plain
                    for table in tables:
                        for row in table:
                            for cell in row:
                                if cell:
                                    used_text = used_text.replace(cell, "")
                    leftover = used_text.strip()
                    if leftover:
                        page_parts.append(leftover)
                pages_text.append("\n\n".join(page_parts))
            else:
                plain = page.extract_text()
                if plain:
                    pages_text.append(f"## Страница {i}\n\n{plain}")

    return "\n\n---\n\n".join(pages_text), has_tables


def _process_docx(path: Path) -> tuple[str, bool]:
    from docx import Document
    doc = Document(str(path))
    parts = []
    has_tables = bool(doc.tables)

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for i, table in enumerate(doc.tables, 1):
        parts.append(f"\n### Таблица {i}")
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        parts.append(_table_to_markdown(rows))

    return "\n\n".join(parts), has_tables


def _process_xlsx(path: Path) -> tuple[str, bool]:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(cell is not None for cell in row):
                rows.append([str(cell) if cell is not None else "" for cell in row])

        if not rows:
            continue

        parts.append(f"## Лист: {sheet_name}")
        parts.append(_table_to_markdown(rows))

    wb.close()
    return "\n\n".join(parts), True


def process_all():
    found_any = False

    for search_dir in SEARCH_DIRS:
        if not search_dir.exists():
            print(f"Папката {search_dir}/ не е намерена — пропускам.")
            continue

        files = [f for f in search_dir.rglob("*")
                 if f.is_file() and f.suffix.lower() in (".pdf", ".docx", ".xlsx")]

        if not files:
            print(f"Няма файлове за обработка в {search_dir}/")
            continue

        print(f"\n📂 {search_dir}/")
        found_any = True

        for path in files:
            txt_out = path.with_suffix(".txt")
            md_out = path.with_suffix(".md")

            if txt_out.exists() or md_out.exists():
                print(f"  Пропускам (вече обработен): {path.name}")
                continue

            print(f"  Обработвам: {path.name} ...", end=" ", flush=True)
            try:
                suffix = path.suffix.lower()
                if suffix == ".pdf":
                    text, has_tables = _process_pdf(path)
                elif suffix == ".docx":
                    text, has_tables = _process_docx(path)
                elif suffix == ".xlsx":
                    text, has_tables = _process_xlsx(path)

                out_path = md_out if has_tables else txt_out
                label = "таблици → .md" if has_tables else "текст → .txt"

                out_path.write_text(text, encoding="utf-8")
                print(f"{label}  ✓")

            except Exception as e:
                print(f"ГРЕШКА: {e}")

    if found_any:
        print("\nГотово.")


if __name__ == "__main__":
    process_all()
